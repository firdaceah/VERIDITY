import fitz
from docx import Document
from io import BytesIO

from analysis.document_pdf_utils import generate_annotated_pdf


def _sample_pdf_bytes():
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        fitz.Point(72, 72),
        "Kalimat ini dibuat oleh AI untuk menguji laporan Veridity.",
        fontsize=12,
    )
    page.insert_text(
        fitz.Point(72, 104),
        "Kalimat kedua ini murni ditulis manusia.",
        fontsize=12,
    )
    payload = doc.write()
    doc.close()
    return payload


def _annotation_count(pdf_bytes):
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    total = 0
    for page in doc:
        annotations = page.annots()
        if annotations:
            total += sum(1 for _ in annotations)
    doc.close()
    return total


def _drawing_count(pdf_bytes):
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    total = sum(len(page.get_drawings()) for page in doc)
    doc.close()
    return total


def _text_content(pdf_bytes):
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    text = "\n".join(page.get_text("text") for page in doc)
    doc.close()
    return text


def _title(pdf_bytes):
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    title = doc.metadata.get("title")
    doc.close()
    return title


def _sample_docx_bytes():
    doc = Document()
    doc.add_paragraph("Kalimat ini dibuat oleh AI untuk menguji laporan Veridity.")
    doc.add_paragraph("Kalimat kedua ini murni ditulis manusia.")
    stream = BytesIO()
    doc.save(stream)
    return stream.getvalue()


def test_generate_annotated_pdf_highlights_non_human_sentence():
    result = generate_annotated_pdf(
        _sample_pdf_bytes(),
        {
            "Kalimat ini dibuat oleh AI untuk menguji laporan Veridity.": "AI-generated",
        },
    )

    assert _annotation_count(result.getvalue()) >= 1
    assert _drawing_count(result.getvalue()) >= 1


def test_generate_annotated_pdf_skips_human_written_sentence():
    result = generate_annotated_pdf(
        _sample_pdf_bytes(),
        {
            "Kalimat kedua ini murni ditulis manusia.": "Human-written",
        },
    )

    assert _annotation_count(result.getvalue()) == 0


def test_generate_annotated_pdf_includes_nlp_metrics_and_consistent_title():
    result = generate_annotated_pdf(
        _sample_pdf_bytes(),
        {
            "Kalimat ini dibuat oleh AI untuk menguji laporan Veridity.": "AI-generated",
        },
        document_metrics={"human_p": 12.5, "ai_p": 75, "hybrid_p": 12.5},
        interpretation="Mayoritas kalimat terindikasi AI.",
    )

    content = _text_content(result.getvalue())
    assert "RINCIAN KOMPUTASI BAHASA" in content
    assert "75.00%" in content
    assert _title(result.getvalue()) == "Laporan Investigasi Forensik Veridity"


def test_generate_annotated_pdf_supports_docx_source_with_highlights():
    result = generate_annotated_pdf(
        _sample_docx_bytes(),
        {
            "Kalimat ini dibuat oleh AI untuk menguji laporan Veridity.": "AI-generated",
        },
        source_extension="docx",
    )

    assert _annotation_count(result.getvalue()) >= 1
    assert "Kalimat ini dibuat oleh AI" in _text_content(result.getvalue())


def test_generate_annotated_pdf_skips_administrative_identity_sentences():
    result = generate_annotated_pdf(
        _sample_pdf_bytes(),
        {
            "Nama Dosen Politeknik Elektronika Negeri Surabaya.": "AI-generated",
            "Kalimat ini dibuat oleh AI untuk menguji laporan Veridity.": "AI-generated",
        },
    )

    doc = fitz.open(stream=result.getvalue(), filetype="pdf")
    subject = doc.metadata.get("subject", "")
    doc.close()
    assert "skipped_administrative" in subject
    assert _annotation_count(result.getvalue()) >= 1
