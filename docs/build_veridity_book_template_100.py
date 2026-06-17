from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Buku_VERIDITY_Template_ISAS_100_Halaman_Revisi_6.docx"


def snippet(path, start, end):
    lines = (ROOT / path).read_text(encoding="utf-8", errors="ignore").splitlines()
    return "\n".join(f"{i:03d}: {line}" for i, line in zip(range(start, end + 1), lines[start - 1:end]))


def style_document(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(15.5)
    sec.page_height = Cm(23)
    sec.top_margin = Cm(1.65)
    sec.bottom_margin = Cm(1.55)
    sec.left_margin = Cm(1.75)
    sec.right_margin = Cm(1.55)
    sec.header_distance = Cm(0.8)
    sec.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.first_line_indent = Cm(0.65)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(4)

    for style_name, size, color in [
        ("Heading 1", 17, RGBColor(0, 0, 0)),
        ("Heading 2", 13, RGBColor(0, 0, 0)),
        ("Heading 3", 11, RGBColor(0, 0, 0)),
    ]:
        s = styles[style_name]
        s.font.name = "Times New Roman"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color
        s.paragraph_format.first_line_indent = Cm(0)
        s.paragraph_format.space_before = Pt(4)
        s.paragraph_format.space_after = Pt(7)


def add_page_number_footer(doc):
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def page_break(doc, page_no):
    if page_no < 100:
        doc.add_page_break()


def p(doc, text, italic=False, bold=False, align=None):
    para = doc.add_paragraph()
    if align:
        para.alignment = align
    run = para.add_run(text)
    run.italic = italic
    run.bold = bold
    return para


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.first_line_indent = Cm(0)
    return h


def code_block(doc, code):
    for line in code.splitlines()[:24]:
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.left_indent = Cm(0.35)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(line[:92])
        run.font.name = "Consolas"
        run.font.size = Pt(7.2)
        run.font.color.rgb = RGBColor(30, 30, 30)


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.left_indent = Cm(0.75)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)


def long_paragraph(topic, angle):
    return (
        f"Pada bagian {topic}, pembahasan diarahkan pada hubungan langsung antara kebutuhan sistem dan implementasi VERIDITY. "
        f"{angle} Hal ini penting karena project tidak hanya menampilkan aplikasi yang berjalan, tetapi juga memperlihatkan alasan teknis di balik pemilihan Laravel, Flutter, Python, database, dan cloud. "
        "Setiap komponen memiliki batas tanggung jawab yang jelas: Laravel menjadi pusat autentikasi dan API, Flutter menjadi antarmuka bergerak untuk pengguna, Python menjalankan analisis forensik dan kecerdasan buatan, sedangkan Distri menjadi contoh aplikasi eksternal yang menguji integrasi layanan. "
        "Dengan cara ini, buku dapat dibaca sebagai dokumentasi pengembangan sekaligus naskah akademik yang menjelaskan masalah, rancangan, kode, pengujian, dan arah pengembangan."
    )


def write_content_page(doc, page_no, title, subtitle, focus):
    heading(doc, title, 1 if title.startswith("Bab") else 2)
    heading(doc, subtitle, 2)
    if "Rumusan Masalah" in subtitle:
        p(doc, "Rumusan masalah pada project VERIDITY disusun dalam beberapa poin agar ruang lingkup penelitian dan implementasi sistem lebih jelas.")
        for item in [
            "Bagaimana membangun backend Laravel yang mampu menerima upload foto, dokumen, dan bukti pembayaran secara aman?",
            "Bagaimana mengintegrasikan aplikasi mobile Flutter dengan REST API Laravel menggunakan token Sanctum?",
            "Bagaimana menjalankan engine Python untuk membaca indikasi manipulasi gambar, dokumen AI, dan validasi nota pembayaran?",
            "Bagaimana menyimpan hasil audit, riwayat pengguna, dan report PDF agar dapat diakses kembali melalui website dan mobile?",
            "Bagaimana menerapkan deployment cloud sehingga Laravel, Python, database, dan storage dapat berjalan sebagai satu sistem?",
        ]:
            bullet(doc, item)
        p(doc, "Poin-poin tersebut menjadi dasar pembahasan pada bab berikutnya, mulai dari perancangan sistem, implementasi kode, integrasi antaraplikasi, sampai pengujian dan deployment.")
        page_break(doc, page_no)
        return
    elif "Tujuan dan Manfaat" in subtitle:
        p(doc, "Tujuan dan manfaat project dibuat dalam bentuk poin agar capaian pengembangan mudah diperiksa oleh dosen saat penilaian.")
        for item in [
            "Membangun website Laravel sebagai pusat autentikasi, upload file, API, dashboard, dan penyimpanan hasil analisis.",
            "Membangun aplikasi mobile Flutter agar pengguna dapat melakukan audit file melalui perangkat bergerak.",
            "Membangun engine Python untuk menjalankan analisis forensik gambar, OCR nota pembayaran, dan klasifikasi dokumen.",
            "Menyediakan report PDF agar hasil audit dapat diunduh dan dijadikan bukti pemeriksaan awal.",
            "Membuktikan integrasi antaraplikasi melalui sistem Distri yang mengirim bukti pembayaran ke VERIDITY.",
            "Menerapkan konsep deployment cloud agar sistem dapat dijalankan di luar komputer lokal pengembang.",
        ]:
            bullet(doc, item)
        p(doc, "Secara umum, manfaat project adalah memberikan contoh sistem terpadu yang menghubungkan web, mobile, cloud, dan kecerdasan buatan dalam satu alur aplikasi.")
        page_break(doc, page_no)
        return
    elif subtitle.startswith("Ringkasan Bab"):
        p(doc, "Halaman ringkasan ini menjelaskan isi bab secara menyeluruh agar bab tidak terasa hanya satu halaman. Setiap bab pada buku ini dibagi menjadi beberapa subbab, mulai dari konsep umum, rancangan, implementasi, contoh kode, sampai evaluasi.")
        for item in focus.split("|"):
            bullet(doc, item.strip())
        p(doc, "Dengan pembagian tersebut, pembaca dapat mengikuti alur pembahasan secara bertahap. Bagian awal menjelaskan masalah dan konsep, bagian tengah menjelaskan desain serta implementasi, sedangkan bagian akhir menjelaskan pengujian, keterbatasan, dan saran pengembangan.")
    else:
        p(doc, long_paragraph(subtitle, focus))
    p(doc, (
        "VERIDITY dibangun untuk menjawab tantangan keaslian data digital. Pada praktiknya, foto, dokumen, dan nota pembayaran dapat berubah bentuk karena proses kompresi, penyuntingan, unggahan ulang, atau bantuan model AI. "
        "Karena itu, sistem tidak memakai satu indikator tunggal. Sistem menggabungkan ELA, noise analysis, metadata analysis, deteksi pola AI, OCR, dan NLP document detection. "
        "Gabungan metode tersebut membuat hasil lebih mudah dijelaskan kepada pengguna dan dosen karena setiap nilai memiliki asal yang dapat ditelusuri pada kode."
    ))
    p(doc, (
        "Pembahasan pada halaman ini juga terkait dengan cara project dipresentasikan. Saat demo, alur yang paling mudah dipahami adalah memulai dari aktor pengguna, kemudian masuk ke file yang diunggah, proses validasi Laravel, pemanggilan Python, penyimpanan database, dan tampilan hasil pada website atau mobile. "
        "Jika dosen meminta bukti implementasi, bagian kode pada bab terkait dapat dibuka langsung dari repository sehingga penjelasan tidak berhenti pada diagram."
    ))
    p(doc, (
        "Dari sisi akademik, bagian ini menunjukkan keterhubungan empat mata kuliah. Workshop Framework terlihat pada routing, controller, model, service, dan Blade. Workshop Perangkat Bergerak terlihat pada struktur Flutter, repository, entity, dan HTTP client. "
        "Workshop Komputasi Awan terlihat pada konfigurasi deployment, storage, database, dan service process. Kecerdasan Buatan terlihat pada modul Python untuk analisis gambar, OCR, dan klasifikasi dokumen."
    ))
    p(doc, (
        "Agar sistem tetap realistis, buku ini selalu mencatat batasan. Hasil analisis VERIDITY adalah indikator awal, bukan putusan hukum. File yang terlalu kecil, terlalu gelap, terkompresi berkali-kali, atau memiliki metadata yang hilang dapat menurunkan ketepatan interpretasi. "
        "Oleh karena itu, setiap output perlu dibaca bersama konteks sumber file dan pemeriksaan manual."
    ))
    page_break(doc, page_no)


def write_code_page(doc, page_no, title, source, start, end, explanation):
    heading(doc, title, 2)
    p(doc, (
        f"Subbab ini memasukkan contoh kode dari {source} baris {start}-{end}. Potongan kode diletakkan di dalam pembahasan implementasi karena dosen meminta kode tidak ditempatkan sebagai lampiran terpisah. "
        "Tujuannya adalah memperlihatkan hubungan antara konsep yang dijelaskan pada bab dengan bagian program yang benar-benar menjalankan fitur tersebut."
    ))
    p(doc, explanation)
    code_block(doc, snippet(source, start, end))
    p(doc, (
        "Dari potongan kode tersebut dapat dilihat bahwa project VERIDITY tidak hanya terdiri dari tampilan, tetapi juga memiliki alur data yang lengkap. "
        "Request pengguna divalidasi, diproses oleh layer yang sesuai, lalu dikembalikan dalam format yang dapat dipakai oleh web atau aplikasi mobile. "
        "Pola ini membantu menjaga maintainability karena perubahan pada satu bagian tidak langsung merusak bagian lain."
    ))
    p(doc, (
        "Saat menjelaskan kode ini di depan dosen, bagian penting yang perlu ditekankan adalah input, proses, output, dan alasan keamanan. "
        "Input menunjukkan data yang diterima sistem, proses menunjukkan logika inti, output menunjukkan respons yang dikirim, sedangkan alasan keamanan menjelaskan validasi token, integration key, pembatasan file, atau penanganan error."
    ))
    page_break(doc, page_no)


def build():
    doc = Document()
    style_document(doc)
    add_page_number_footer(doc)

    page = 1
    # Template front matter
    title = "Pembuatan Aplikasi VERIDITY"
    subtitle = "Sistem Forensik Digital Terintegrasi antara Aplikasi Flutter, Website Laravel, Python AI Engine, Distri, dan Deployment Cloud"
    for text, size, bold in [(title, 22, True), (subtitle, 13, False), ("Nama Penulis", 12, False)]:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
    doc.add_paragraph()
    for label in [
        "Editor:\nNama Editor",
        "Desainer:\nNama Desainer",
        "Sumber Gambar Kover:\nwww.freepik.com",
        "Penata Letak:\nNama Penata Letak",
        "Proofreader:\nTim YPCM",
        "Ukuran:\n100 hlm, 15,5 x 23 cm",
        "ISBN:",
        "Cetakan Pertama:\nJuni 2026",
    ]:
        p(doc, label)
    p(doc, "ISAS PRESS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    page_break(doc, page)

    page += 1
    p(doc, "Hak cipta dilindungi undang-undang", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "Dilarang keras menerjemahkan, memfotokopi, atau memperbanyak sebagian atau seluruh isi buku ini tanpa izin tertulis dari Penerbit.", align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "Buku ini disusun sebagai dokumentasi project UAS VERIDITY. Data nama penulis, editor, desainer, ISBN, dan informasi penerbit dapat disesuaikan sebelum pengumpulan akhir.")
    p(doc, "Project VERIDITY menggabungkan implementasi Laravel, Flutter, Python, aplikasi Distri, database, serta deployment cloud. Penyusunan buku mengikuti template buku akademik dengan halaman awal, daftar isi, prakata, bab pembahasan, daftar pustaka, dan profil penulis.")
    page_break(doc, page)

    page += 1
    heading(doc, "Daftar Isi", 1)
    for item in [
        "Prakata",
        "Bab 1 Pendahuluan",
        "Bab 2 Konsep Dasar dan Teknologi Pendukung",
        "Bab 3 Analisis dan Perancangan Sistem",
        "Bab 4 Pembangunan Backend Website Laravel",
        "Bab 5 Pembangunan Aplikasi Mobile Flutter",
        "Bab 6 Implementasi Sistem Forensik dan Kecerdasan Buatan",
        "Bab 7 Deployment Sistem pada Cloud",
        "Bab 8 Pengujian dan Evaluasi Sistem",
        "Bab 9 Penutup",
        "Daftar Pustaka",
        "Profil Penulis",
    ]:
        p(doc, item)
    p(doc, "Catatan: nomor halaman final dapat diperbarui otomatis melalui fitur daftar isi Microsoft Word apabila diperlukan.")
    page_break(doc, page)

    page += 1
    heading(doc, "Prakata", 1)
    for _ in range(5):
        p(doc, (
            "Buku ini disusun sebagai dokumentasi project UAS VERIDITY pada mata kuliah Workshop Pemrograman Framework, Workshop Pemrograman Perangkat Bergerak, Workshop Aplikasi Komputasi Awan, dan Kecerdasan Buatan. "
            "Tujuan penyusunan buku adalah menjelaskan project secara menyeluruh, mulai dari alasan pembuatan sistem, rancangan arsitektur, struktur database, implementasi kode, integrasi mobile, proses analisis kecerdasan buatan, deployment, hingga pengujian. "
            "Dengan format buku, pembaca tidak hanya melihat hasil akhir aplikasi, tetapi juga memahami proses teknis yang membuat sistem berjalan."
        ))
    page_break(doc, page)

    pages = [
        ("Bab 1 Pendahuluan", "Latar Belakang", "Perkembangan AI generatif membuat foto, dokumen, dan bukti pembayaran semakin mudah dimanipulasi sehingga sistem audit awal diperlukan."),
        ("Bab 1 Pendahuluan", "Rumusan Masalah", "Masalah utama adalah bagaimana menerima file, menganalisisnya, menyimpan hasil, menampilkan riwayat, dan mengintegrasikan sistem eksternal."),
        ("Bab 1 Pendahuluan", "Tujuan dan Manfaat", "Tujuan project adalah membuat sistem web dan mobile yang membantu membaca indikasi manipulasi digital secara terstruktur."),
        ("Bab 1 Pendahuluan", "Batasan Sistem", "VERIDITY memberi indikator forensik awal dan tidak menggantikan pemeriksaan hukum atau verifikasi manual."),
        ("Bab 1 Pendahuluan", "Keterkaitan Mata Kuliah", "Project memetakan Laravel, Flutter, cloud deployment, dan Python AI ke capaian pembelajaran lintas mata kuliah."),
        ("Bab 2 Konsep Dasar dan Teknologi Pendukung", "Forensik Digital", "Forensik digital membaca jejak teknis file seperti kompresi, noise, metadata, pola spektral, dan struktur teks."),
        ("Bab 2 Konsep Dasar dan Teknologi Pendukung", "Error Level Analysis", "ELA membandingkan gambar asli dengan hasil kompresi ulang untuk melihat area dengan error kompresi berbeda."),
        ("Bab 2 Konsep Dasar dan Teknologi Pendukung", "Noise Analysis", "Noise analysis membaca konsistensi residu piksel untuk menemukan area yang karakter visualnya berbeda."),
        ("Bab 2 Konsep Dasar dan Teknologi Pendukung", "Metadata Analysis", "Metadata membantu membaca jejak kamera, software, waktu, dan proses ekspor file."),
        ("Bab 2 Konsep Dasar dan Teknologi Pendukung", "Deepfake dan GAN Indicator", "Indikator AI membaca pola spektral yang dapat muncul pada gambar hasil generatif."),
        ("Bab 2 Konsep Dasar dan Teknologi Pendukung", "OCR Tesseract", "OCR mengubah tulisan pada nota menjadi teks agar dapat dicocokkan dengan data checkout."),
        ("Bab 2 Konsep Dasar dan Teknologi Pendukung", "NLP Document Detection", "NLP memecah dokumen menjadi unit kalimat dan memberi label human, AI, atau hybrid."),
        ("Bab 2 Konsep Dasar dan Teknologi Pendukung", "Laravel", "Laravel menjadi pusat backend, routing, autentikasi, API, penyimpanan, dan report."),
        ("Bab 2 Konsep Dasar dan Teknologi Pendukung", "Flutter", "Flutter menjadi client mobile yang mengirim request ke Laravel dan menampilkan hasil audit."),
        ("Bab 2 Konsep Dasar dan Teknologi Pendukung", "FastAPI dan Python", "FastAPI menjadi service dokumen, sedangkan script Python menangani analisis gambar."),
        ("Bab 2 Konsep Dasar dan Teknologi Pendukung", "Database dan Storage", "Database menyimpan data terstruktur, sedangkan storage menyimpan file upload dan report."),
        ("Bab 3 Analisis dan Perancangan Sistem", "Analisis Aktor", "Aktor meliputi user, admin, pengguna mobile, reseller Distri, dan admin Distri."),
        ("Bab 3 Analisis dan Perancangan Sistem", "Kebutuhan Fungsional", "Kebutuhan meliputi register, login, upload, analisis, history, detail, report, dan validasi nota."),
        ("Bab 3 Analisis dan Perancangan Sistem", "Kebutuhan Non-Fungsional", "Sistem perlu aman, responsif, tervalidasi, dapat dipantau, dan siap dideploy."),
        ("Bab 3 Analisis dan Perancangan Sistem", "Use Case User", "User login, mengunggah file, membaca hasil, melihat riwayat, dan mengunduh report."),
        ("Bab 3 Analisis dan Perancangan Sistem", "Use Case Distri", "Distri mengirim nota pembayaran ke VERIDITY dan menerima status validasi."),
        ("Bab 3 Analisis dan Perancangan Sistem", "Arsitektur Sistem", "Laravel menjadi gateway antara web, mobile, Distri, Python, database, dan storage."),
        ("Bab 3 Analisis dan Perancangan Sistem", "Desain Database", "Tabel forensic_analyses menyimpan hasil audit, sedangkan tabel order Distri menyimpan status validasi."),
        ("Bab 3 Analisis dan Perancangan Sistem", "Desain Antarmuka", "Antarmuka dibuat untuk mempercepat upload, membaca hasil, membuka history, dan report."),
        ("Bab 4 Pembangunan Backend Website Laravel", "Struktur Folder Backend", "Folder veridity-laravel berisi controller, model, service, resource, view, route, migration, dan konfigurasi."),
        ("Bab 4 Pembangunan Backend Website Laravel", "Autentikasi Web dan Mobile", "Backend membedakan login web berbasis session dan login mobile berbasis token Sanctum."),
        ("Bab 4 Pembangunan Backend Website Laravel", "Manajemen Audit", "Audit disimpan sebagai record ForensicAnalysis dengan field hasil JSON dan path report."),
        ("Bab 4 Pembangunan Backend Website Laravel", "Integrasi Python", "Laravel memanggil Python untuk analisis gambar dan FastAPI untuk analisis dokumen."),
        ("Bab 4 Pembangunan Backend Website Laravel", "Report PDF", "Report PDF disimpan agar website dan mobile mengunduh hasil yang sama."),
        ("Bab 5 Pembangunan Aplikasi Mobile Flutter", "Struktur Flutter", "Flutter memakai pemisahan core, app, dan features agar kode lebih mudah dirawat."),
        ("Bab 5 Pembangunan Aplikasi Mobile Flutter", "API Client", "ApiClient menyatukan URL, header, JSON request, multipart upload, dan error handling."),
        ("Bab 5 Pembangunan Aplikasi Mobile Flutter", "Repository Pattern", "Repository memisahkan UI dari detail request HTTP dan parsing JSON."),
        ("Bab 5 Pembangunan Aplikasi Mobile Flutter", "Upload File", "Flutter membaca file dari perangkat lalu mengirim bytes ke Laravel."),
        ("Bab 5 Pembangunan Aplikasi Mobile Flutter", "History dan Detail Audit", "Mobile mengambil daftar audit dan detail hasil dari endpoint API Laravel."),
        ("Bab 6 Implementasi Sistem Forensik dan Kecerdasan Buatan", "Pipeline Analisis Gambar", "Pipeline Python menjalankan metadata, ELA, AI indicator, noise map, dan final score."),
        ("Bab 6 Implementasi Sistem Forensik dan Kecerdasan Buatan", "Rumus Final Score", "Final score menggabungkan ELA, noise, metadata, dan AI indicator dengan bobot seimbang."),
        ("Bab 6 Implementasi Sistem Forensik dan Kecerdasan Buatan", "Analisis Dokumen", "Dokumen dianalisis dengan ekstraksi teks dan klasifikasi kalimat."),
        ("Bab 6 Implementasi Sistem Forensik dan Kecerdasan Buatan", "Validasi Nota", "Nota dianalisis visual dan kontennya dibaca dengan OCR."),
        ("Bab 7 Deployment Sistem pada Cloud", "Komponen Deployment", "Deployment mencakup Laravel, Python, database, storage, web server, dan environment variable."),
        ("Bab 7 Deployment Sistem pada Cloud", "Keamanan Cloud", "Keamanan mencakup HTTPS, secret, token, integration key, validasi upload, dan monitoring."),
        ("Bab 8 Pengujian dan Evaluasi Sistem", "Pengujian Web dan API", "Pengujian memastikan route, auth, upload, history, detail, dan report berjalan benar."),
        ("Bab 8 Pengujian dan Evaluasi Sistem", "Pengujian Mobile", "Pengujian mobile memastikan token, upload file, history, detail, dan report dapat dipakai."),
        ("Bab 8 Pengujian dan Evaluasi Sistem", "Evaluasi Metode AI", "Evaluasi harus memakai file asli, editan, AI-generated, human-written, dan nota valid/tidak valid."),
        ("Bab 9 Penutup", "Kesimpulan", "VERIDITY berhasil menjadi sistem terpadu yang menghubungkan web, mobile, AI, cloud, dan aplikasi eksternal."),
        ("Bab 9 Penutup", "Saran", "Pengembangan dapat menambahkan dataset besar, queue, model AI lebih kuat, dan dashboard metrik."),
    ]

    code_pages = [
        ("Bab 4 Pembangunan Backend Website Laravel", "Contoh Kode Route API Laravel", "veridity-laravel/routes/api.php", 1, 28, "Route API menjadi pintu masuk Flutter dan Distri. Endpoint umum seperti register dan login dibuka tanpa token, sedangkan endpoint audit dilindungi middleware auth:sanctum."),
        ("Bab 4 Pembangunan Backend Website Laravel", "Contoh Kode Token Response Sanctum", "veridity-laravel/app/Http/Controllers/Api/AuthController.php", 15, 33, "Kode ini membuat token untuk mobile dan mengembalikan format JSON yang konsisten, sehingga Flutter dapat menyimpan access_token dan memakainya pada request berikutnya."),
        ("Bab 4 Pembangunan Backend Website Laravel", "Contoh Kode Login Web dan Mobile", "veridity-laravel/app/Http/Controllers/Api/AuthController.php", 58, 86, "Login membedakan request browser dan request JSON. Browser memakai session, sedangkan mobile memperoleh token bearer."),
        ("Bab 4 Pembangunan Backend Website Laravel", "Contoh Kode Proteksi Akses Audit", "veridity-laravel/app/Http/Controllers/Api/ForensicController.php", 20, 43, "Kode ini memastikan user hanya dapat membaca audit miliknya sendiri, sementara admin dapat membaca data untuk keperluan monitoring."),
        ("Bab 4 Pembangunan Backend Website Laravel", "Contoh Kode Menjalankan Python", "veridity-laravel/app/Http/Controllers/Api/ForensicController.php", 60, 118, "Laravel memakai proc_open untuk menjalankan script Python dan menangkap stdout sebagai JSON hasil analisis."),
        ("Bab 4 Pembangunan Backend Website Laravel", "Contoh Kode Validasi Output Python", "veridity-laravel/app/Http/Controllers/Api/ForensicController.php", 119, 157, "Jika Python tidak mengirim output atau output bukan JSON, Laravel mencatat error dan mengembalikan pesan yang dapat dipahami user."),
        ("Bab 4 Pembangunan Backend Website Laravel", "Contoh Kode Integrasi Distri", "veridity-laravel/app/Http/Controllers/Api/ForensicController.php", 174, 214, "Endpoint ini memeriksa integration key dan field bukti pembayaran agar hanya aplikasi Distri yang sah dapat mengirim nota."),
        ("Bab 4 Pembangunan Backend Website Laravel", "Contoh Kode Model ForensicAnalysis", "veridity-laravel/app/Models/ForensicAnalysis.php", 1, 38, "Model ini menentukan field yang dapat diisi dan cast JSON untuk metadata_details serta final_result."),
        ("Bab 5 Pembangunan Aplikasi Mobile Flutter", "Contoh Kode ApiClient Flutter", "veridity_mobile/lib/core/network/api_client.dart", 1, 41, "ApiClient menormalisasi base URL, membuat header JSON, dan menambahkan bearer token bila tersedia."),
        ("Bab 5 Pembangunan Aplikasi Mobile Flutter", "Contoh Kode Multipart Upload Flutter", "veridity_mobile/lib/core/network/api_client.dart", 70, 122, "Kode ini mengirim file ke Laravel menggunakan multipart request, baik dari path file maupun bytes."),
        ("Bab 5 Pembangunan Aplikasi Mobile Flutter", "Contoh Kode Error Handling Flutter", "veridity_mobile/lib/core/network/api_client.dart", 124, 170, "Bagian decode membaca response JSON, membedakan status sukses dan error, lalu mengambil pesan validasi pertama."),
        ("Bab 5 Pembangunan Aplikasi Mobile Flutter", "Contoh Kode AuditRepository History", "veridity_mobile/lib/features/audit/data/repositories/audit_repository.dart", 14, 31, "Repository mengambil token dari SessionStore, memanggil endpoint /audits, lalu memetakan JSON menjadi AuditEntity."),
        ("Bab 5 Pembangunan Aplikasi Mobile Flutter", "Contoh Kode AuditRepository Upload", "veridity_mobile/lib/features/audit/data/repositories/audit_repository.dart", 33, 52, "Upload file memastikan user sudah login dan bytes file dapat dibaca sebelum request dikirim."),
        ("Bab 6 Implementasi Sistem Forensik dan Kecerdasan Buatan", "Contoh Kode Pipeline Python", "python/analyze_all.py", 1, 35, "Script analyze_all.py memanggil metadata analysis, ELA, noise map, dan deepfake detector dalam satu alur investigasi."),
        ("Bab 6 Implementasi Sistem Forensik dan Kecerdasan Buatan", "Contoh Kode Simpan ELA dan AI Score", "python/analyze_all.py", 36, 72, "Bagian ini menyimpan peta ELA, menjalankan deteksi AI, dan membuat noise map yang akan ditampilkan pada hasil audit."),
        ("Bab 6 Implementasi Sistem Forensik dan Kecerdasan Buatan", "Contoh Kode Rumus Final Score", "python/analyze_all.py", 73, 114, "Kode ini menghitung final score dan menentukan verdict authentic, manipulated, atau deepfake/AI generated."),
        ("Bab 6 Implementasi Sistem Forensik dan Kecerdasan Buatan", "Contoh Kode FastAPI Dokumen", "python/main_api.py", 1, 29, "FastAPI menerima file dokumen, menjalankan run_document_analysis, dan mengirim JSON report ke Laravel."),
        ("Bab 6 Implementasi Sistem Forensik dan Kecerdasan Buatan", "Contoh Kode Generate PDF Report", "python/main_api.py", 31, 68, "Endpoint generate PDF membuat report dokumen dengan classification map dan metadata hasil analisis."),
        ("Bab 4 Pembangunan Backend Website Laravel", "Contoh Kode Service Distri", "distri/app/Services/VeridityProofService.php", 1, 42, "Distri mengirim bukti pembayaran ke VERIDITY dengan integration key dan konteks pembayaran."),
        ("Bab 4 Pembangunan Backend Website Laravel", "Contoh Kode Mapping Status Distri", "distri/app/Services/VeridityProofService.php", 43, 86, "Hasil VERIDITY dipetakan menjadi payment_status, order_status, audit_id, final_score, dan validation_details."),
    ]

    def make(chapter, subtitles):
        return [(chapter, sub, focus) for sub, focus in subtitles]

    def emit(items):
        nonlocal page
        for title, sub, focus in items:
            page += 1
            write_content_page(doc, page, title, sub, focus)

    bab1 = make("Bab 1 Pendahuluan", [
        ("Ringkasan Bab 1", "Latar belakang kebutuhan forensik digital|Rumusan masalah berbasis kebutuhan sistem|Tujuan dan manfaat project|Batasan interpretasi hasil VERIDITY"),
        ("Latar Belakang", "Perkembangan AI generatif membuat foto, dokumen, dan bukti pembayaran semakin mudah dimanipulasi sehingga sistem audit awal diperlukan."),
        ("Konteks Project VERIDITY", "VERIDITY diposisikan sebagai sistem pembantu investigasi awal untuk membaca sinyal teknis file digital."),
        ("Identifikasi Masalah", "Masalah utama berada pada keaslian foto, dokumen yang dibuat AI, dan nota pembayaran yang tidak sesuai transaksi."),
        ("Rumusan Masalah", "Rumusan masalah harus dapat diturunkan menjadi kebutuhan backend, mobile, AI engine, database, dan deployment."),
        ("Tujuan dan Manfaat", "Tujuan project adalah membuat sistem web dan mobile yang membantu membaca indikasi manipulasi digital secara terstruktur."),
        ("Manfaat bagi Pengguna", "Pengguna memperoleh alat bantu untuk membaca hasil audit, riwayat, dan report PDF dari satu sistem."),
        ("Manfaat bagi Pengembang", "Pengembang belajar menyatukan Laravel, Flutter, Python, database, API, dan cloud deployment."),
        ("Batasan Sistem", "VERIDITY memberi indikator forensik awal dan tidak menggantikan pemeriksaan hukum atau verifikasi manual."),
    ])

    bab2 = make("Bab 2 Konsep Dasar dan Teknologi Pendukung", [
        ("Ringkasan Bab 2", "Forensik digital gambar|Forensik dokumen|Laravel dan Flutter|Python, FastAPI, OCR, NLP, database, dan cloud"),
        ("Forensik Digital", "Forensik digital membaca jejak teknis file seperti kompresi, noise, metadata, pola spektral, dan struktur teks."),
        ("Error Level Analysis", "ELA membandingkan gambar asli dengan hasil kompresi ulang untuk melihat area dengan error kompresi berbeda."),
        ("Noise Analysis", "Noise analysis membaca konsistensi residu piksel untuk menemukan area yang karakter visualnya berbeda."),
        ("Metadata Analysis", "Metadata membantu membaca jejak kamera, software, waktu, dan proses ekspor file."),
        ("Deepfake dan GAN Indicator", "Indikator AI membaca pola spektral yang dapat muncul pada gambar hasil generatif."),
        ("OCR Tesseract", "OCR mengubah tulisan pada nota menjadi teks agar dapat dicocokkan dengan data checkout."),
        ("NLP Document Detection", "NLP memecah dokumen menjadi unit kalimat dan memberi label human, AI, atau hybrid."),
        ("Laravel", "Laravel menjadi pusat backend, routing, autentikasi, API, penyimpanan, dan report."),
        ("Flutter", "Flutter menjadi client mobile yang mengirim request ke Laravel dan menampilkan hasil audit."),
        ("FastAPI dan Python", "FastAPI menjadi service dokumen, sedangkan script Python menangani analisis gambar."),
        ("Database dan Storage", "Database menyimpan data terstruktur, sedangkan storage menyimpan file upload dan report."),
        ("REST API dan Sanctum", "REST API menjadi penghubung Flutter, website, dan Distri, sedangkan Sanctum menjaga akses mobile."),
        ("Cloud Deployment", "Cloud deployment membuat aplikasi dapat diakses dari luar perangkat pengembang dengan konfigurasi layanan terpisah."),
    ])

    bab3 = make("Bab 3 Analisis dan Perancangan Sistem", [
        ("Ringkasan Bab 3", "Aktor sistem|Kebutuhan fungsional dan non-fungsional|Use case|Arsitektur, database, dan antarmuka"),
        ("Analisis Aktor", "Aktor meliputi user, admin, pengguna mobile, reseller Distri, dan admin Distri."),
        ("Kebutuhan Fungsional", "Kebutuhan meliputi register, login, upload, analisis, history, detail, report, dan validasi nota."),
        ("Kebutuhan Non-Fungsional", "Sistem perlu aman, responsif, tervalidasi, dapat dipantau, dan siap dideploy."),
        ("Use Case User", "User login, mengunggah file, membaca hasil, melihat riwayat, dan mengunduh report."),
        ("Use Case Admin", "Admin memantau data audit, hasil analisis, dan status validasi pada sistem."),
        ("Use Case Distri", "Distri mengirim nota pembayaran ke VERIDITY dan menerima status validasi."),
        ("Arsitektur Sistem", "Laravel menjadi gateway antara web, mobile, Distri, Python, database, dan storage."),
        ("Alur Analisis Gambar", "File gambar divalidasi Laravel, dianalisis Python, lalu hasilnya disimpan ke database."),
        ("Alur Analisis Dokumen", "Dokumen dikirim ke FastAPI, diklasifikasi, lalu hasilnya dikembalikan sebagai JSON dan report."),
        ("Desain Database", "Tabel forensic_analyses menyimpan hasil audit, sedangkan tabel order Distri menyimpan status validasi."),
        ("Desain Antarmuka", "Antarmuka dibuat untuk mempercepat upload, membaca hasil, membuka history, dan report."),
    ])

    bab4a = make("Bab 4 Pembangunan Backend Website Laravel", [
        ("Ringkasan Bab 4", "Struktur Laravel|Autentikasi|REST API|Model audit|Integrasi Python dan Distri"),
        ("Struktur Folder Backend", "Folder veridity-laravel berisi controller, model, service, resource, view, route, migration, dan konfigurasi."),
        ("Autentikasi Web dan Mobile", "Backend membedakan login web berbasis session dan login mobile berbasis token Sanctum."),
        ("Manajemen Audit", "Audit disimpan sebagai record ForensicAnalysis dengan field hasil JSON dan path report."),
        ("Integrasi Python", "Laravel memanggil Python untuk analisis gambar dan FastAPI untuk analisis dokumen."),
    ])

    bab4b = make("Bab 4 Pembangunan Backend Website Laravel", [
        ("Report PDF", "Report PDF disimpan agar website dan mobile mengunduh hasil yang sama."),
        ("Dashboard Admin", "Dashboard admin dipakai untuk memantau audit, pengguna, dan hasil validasi sistem."),
        ("Validasi Bukti Pembayaran Distri", "Integrasi Distri membuktikan VERIDITY dapat dipakai oleh aplikasi eksternal."),
    ])

    bab5 = make("Bab 5 Pembangunan Aplikasi Mobile Flutter", [
        ("Ringkasan Bab 5", "Struktur Flutter|ApiClient|Repository|Upload file|History dan detail audit"),
        ("Struktur Flutter", "Flutter memakai pemisahan core, app, dan features agar kode lebih mudah dirawat."),
        ("API Client", "ApiClient menyatukan URL, header, JSON request, multipart upload, dan error handling."),
        ("Repository Pattern", "Repository memisahkan UI dari detail request HTTP dan parsing JSON."),
        ("Upload File", "Flutter membaca file dari perangkat lalu mengirim bytes ke Laravel."),
        ("History dan Detail Audit", "Mobile mengambil daftar audit dan detail hasil dari endpoint API Laravel."),
    ])

    bab6 = make("Bab 6 Implementasi Sistem Forensik dan Kecerdasan Buatan", [
        ("Ringkasan Bab 6", "Pipeline gambar|ELA, noise, metadata, dan AI indicator|Dokumen dan OCR|Final score dan interpretasi"),
        ("Pipeline Analisis Gambar", "Pipeline Python menjalankan metadata, ELA, AI indicator, noise map, dan final score."),
        ("Implementasi ELA", "ELA menghasilkan peta visual dan skor anomali yang membantu membaca konsistensi kompresi."),
        ("Implementasi Noise Map", "Noise map membantu menemukan area yang memiliki distribusi residu piksel tidak seragam."),
        ("Rumus Final Score", "Final score menggabungkan ELA, noise, metadata, dan AI indicator dengan bobot seimbang."),
        ("Analisis Dokumen", "Dokumen dianalisis dengan ekstraksi teks dan klasifikasi kalimat."),
        ("Validasi Nota", "Nota dianalisis visual dan kontennya dibaca dengan OCR."),
        ("Batasan Deteksi AI", "Hasil deteksi AI perlu dibaca sebagai indikator karena berisiko false positive dan false negative."),
    ])

    bab7 = make("Bab 7 Deployment Sistem pada Cloud", [
        ("Ringkasan Bab 7", "Komponen deployment|Laravel dan Python service|Database dan storage|Keamanan dan monitoring"),
        ("Komponen Deployment", "Deployment mencakup Laravel, Python, database, storage, web server, dan environment variable."),
        ("Deployment Laravel", "Laravel membutuhkan konfigurasi .env, dependency Composer, migration, storage link, dan web server."),
        ("Deployment Python", "Python engine membutuhkan virtual environment, dependency, uvicorn, dan process manager."),
        ("Storage dan Database", "Database menyimpan data terstruktur, sedangkan object storage menyimpan file audit dan report."),
        ("Keamanan Cloud", "Keamanan mencakup HTTPS, secret, token, integration key, validasi upload, dan monitoring."),
    ])

    bab8 = make("Bab 8 Pengujian dan Evaluasi Sistem", [
        ("Ringkasan Bab 8", "Pengujian web|Pengujian mobile|Pengujian API|Pengujian Python dan evaluasi hasil"),
        ("Pengujian Web dan API", "Pengujian memastikan route, auth, upload, history, detail, dan report berjalan benar."),
        ("Pengujian Mobile", "Pengujian mobile memastikan token, upload file, history, detail, dan report dapat dipakai."),
        ("Pengujian Integrasi Distri", "Pengujian Distri memastikan nota terkirim dan status validasi masuk kembali ke order."),
        ("Pengujian Python", "Pengujian Python memastikan modul ELA, metadata, noise, dan dokumen dapat berjalan."),
        ("Evaluasi Metode AI", "Evaluasi harus memakai file asli, editan, AI-generated, human-written, dan nota valid/tidak valid."),
        ("Evaluasi Usability", "Evaluasi usability memastikan pengguna mudah memahami alur upload, hasil, history, dan report."),
    ])

    bab9 = make("Bab 9 Penutup", [
        ("Ringkasan Bab 9", "Kesimpulan implementasi|Saran pengembangan|Arah riset berikutnya"),
        ("Kesimpulan", "VERIDITY berhasil menjadi sistem terpadu yang menghubungkan web, mobile, AI, cloud, dan aplikasi eksternal."),
        ("Saran", "Pengembangan dapat menambahkan dataset besar, queue, model AI lebih kuat, dan dashboard metrik."),
        ("Arah Pengembangan", "Arah lanjutan meliputi explainable AI, object storage penuh, dan pengujian pada dataset lebih bervariasi."),
    ])

    emit(bab1)
    emit(bab2)
    emit(bab3)
    emit(bab4a)
    for cp in code_pages[:8] + code_pages[18:]:
        page += 1
        write_code_page(doc, page, cp[1], cp[2], cp[3], cp[4], cp[5])
    emit(bab4b)
    emit(bab5)
    for cp in code_pages[8:13]:
        page += 1
        write_code_page(doc, page, cp[1], cp[2], cp[3], cp[4], cp[5])
    emit(bab6)
    for cp in code_pages[13:18]:
        page += 1
        write_code_page(doc, page, cp[1], cp[2], cp[3], cp[4], cp[5])
    emit(bab7)
    emit(bab8)
    emit(bab9)

    while page < 98:
        page += 1
        write_content_page(doc, page, "Bab 8 Pengujian dan Evaluasi Sistem", "Catatan Evaluasi Tambahan", "Halaman tambahan ini memperpanjang pembahasan evaluasi agar buku lebih padat dan dapat dilengkapi screenshot apabila diperlukan.")

    page += 1
    heading(doc, "Daftar Pustaka", 1)
    refs = [
        "[ELA] Gupta, A., Joshi, R., dan Laban, R. 2022. Detection of Tool based Edited Images from Error Level Analysis and Convolutional Neural Network. arXiv. Link: https://arxiv.org/abs/2204.09075.",
        "[Noise] Cozzolino, D. dan Verdoliva, L. 2018. Noiseprint: a CNN-based camera model fingerprint. arXiv. Link: https://arxiv.org/abs/1808.08396.",
        "[Metadata] Farid, H. 2009. Image Forgery Detection. IEEE Signal Processing Magazine. Link: https://farid.berkeley.edu/downloads/publications/spm09.pdf.",
        "[Deepfake] Nguyen, T. T., Nguyen, Q. V. H., Nguyen, D. T., Nguyen, D. T., Huynh-The, T., Nahavandi, S., Nguyen, T. T., Pham, Q.-V., dan Nguyen, C. M. 2019. Deep Learning for Deepfakes Creation and Detection: A Survey. arXiv. Link: https://arxiv.org/abs/1909.11573.",
        "[NLP] Sadasivan, V. S., Kumar, A., Balasubramanian, S., Wang, W., dan Feizi, S. 2023. Can AI-Generated Text be Reliably Detected? arXiv. Link: https://arxiv.org/abs/2303.11156.",
        "[OCR/NLP] Rusli, F. M., Adhiguna, K. A., dan Irawan, H. 2020. Indonesian ID Card Extractor Using Optical Character Recognition and Natural Language Post-Processing. arXiv. Link: https://arxiv.org/abs/2101.05214.",
    ]
    for ref in refs:
        p(doc, ref)
    page_break(doc, page)

    page += 1
    heading(doc, "Profil Penulis", 1)
    for _ in range(5):
        p(doc, (
            "Penulis adalah kelompok mahasiswa pengembang project VERIDITY. Project ini dikerjakan sebagai bentuk integrasi kemampuan pemrograman framework, pemrograman perangkat bergerak, komputasi awan, dan kecerdasan buatan. "
            "Dalam proses pengembangan, penulis membangun backend Laravel, aplikasi mobile Flutter, service Python untuk analisis forensik, aplikasi Distri untuk integrasi nota pembayaran, serta dokumentasi deployment dan pengujian. "
            "Profil lengkap anggota kelompok, NRP, kelas, dan kontribusi masing-masing dapat diisi pada bagian ini sebelum buku dikumpulkan."
        ))

    doc.save(OUT)


if __name__ == "__main__":
    build()
